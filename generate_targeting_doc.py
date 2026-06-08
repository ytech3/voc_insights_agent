"""
Generate a 2-3 page Rays-branded Word document:
2026 VoC Insights — Audience Targeting & Messaging to Drive Ticket Page Visits
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# --- BRANDING ---
NAVY = RGBColor(0x09, 0x2C, 0x5C)
SKY = RGBColor(0x8F, 0xBC, 0xE6)
YELLOW = RGBColor(0xF5, 0xD1, 0x30)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)

OUT = r"C:\Users\ytaketani\Music\ytaketani\voc_insights_agent\2026_VoC_Ticket_Page_Targeting_Insights_v2.docx"

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# --- Helper functions ---
def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a branded table with navy headers."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "092C5C")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.color.rgb = DARK_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Alternate row shading
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table

def add_heading_styled(doc, text, level=1):
    """Add heading with Rays navy color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = NAVY
    return heading

# ============================================================
# PAGE 1 — TITLE & EXECUTIVE SUMMARY
# ============================================================

# Title
title = doc.add_heading("", level=0)
run = title.add_run("2026 VoC Insights")
run.font.color.rgb = NAVY
run.font.size = Pt(24)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Audience Targeting & Messaging Strategy to Drive Ticket Page Visits")
run.font.size = Pt(13)
run.font.color.rgb = SKY
run.bold = True

# Separator
doc.add_paragraph("_" * 80).runs[0].font.color.rgb = SKY

# Context
add_heading_styled(doc, "Strategic Context", level=2)
ctx = doc.add_paragraph()
ctx.paragraph_format.space_after = Pt(6)
ctx_text = (
    "In 2026, 72% of surveyed Rays fans (8,332 of 11,561) rated their overall experience a 9 or 10 out of 10. "
    "These highly satisfied fans represent our strongest conversion opportunity: they already love the product. "
    "The question is—what messaging and incentives will bring them back to the ticket purchase page?\n\n"
    "This brief distills VoC survey data into actionable targeting recommendations. "
    "Each segment below identifies WHO to target, WHAT message to lead with, and WHY it resonates based on their stated motivators."
)
run = ctx.add_run(ctx_text)
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY

# Core Insight Box
add_heading_styled(doc, "The #1 Insight", level=2)
insight = doc.add_paragraph()
run = insight.add_run("Ticket discounts and value-based offers are the primary driver across ALL segments (36% rank it #1). ")
run.font.size = Pt(10)
run.bold = True
run.font.color.rgb = NAVY
run2 = insight.add_run(
    "But the secondary hooks differ dramatically by audience. Tailoring the secondary message to segment-specific motivators "
    "is where incremental ticket page visits will come from."
)
run2.font.size = Pt(10)
run2.font.color.rgb = DARK_GRAY

# ============================================================
# PAGE 2 — TARGETING PLAYBOOK TABLE
# ============================================================

add_heading_styled(doc, "Target Audience Playbook", level=1)
intro = doc.add_paragraph()
run = intro.add_run("Priority segments ranked by reach and differentiation of messaging opportunity:")
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY
run.italic = True

headers = ["Target Segment", "Reach", "Primary Driver", "Recommended Message Hook", "Tactic"]

rows = [
    [
        "Women (all ages)",
        "27% of fans",
        "Ticket discounts (42%)\nSafety & comfort",
        "\"Your safe, comfortable night out — at a great price\"",
        "Email campaigns emphasizing clean stadium, comfortable seating, and discount CTA"
    ],
    [
        "Young Adults (25-34)",
        "5.5% of fans",
        "Ticket discounts (42%)\nGiveaways (12%)",
        "\"Exclusive collectibles + unbeatable deals\"",
        "Social media (IG/TikTok) with giveaway visuals + promo code link to ticket page"
    ],
    [
        "Hispanic/Latino Fans",
        "8.7% of fans",
        "Ticket discounts (44%)\nGiveaways (9%)",
        "\"Special offers + exclusive collectibles for you\"",
        "Bilingual email/social; highlight heritage nights + discounts in same CTA"
    ],
    [
        "Black/African American Fans",
        "2.4% of fans",
        "Theme nights (14%)\nParking credits (14%)",
        "\"Theme nights made for you — free parking included\"",
        "Targeted ads for theme night events; bundle parking credit with ticket link"
    ],
    [
        "Solo Attendees",
        "6.6% of fans",
        "Love of baseball (44%)\nTicket discounts (20%)",
        "\"Your game. Your time. Affordable single seats available.\"",
        "Single-ticket promos via app push notifications; emphasize staff/atmosphere"
    ],
    [
        "Friends Groups",
        "22% of fans",
        "Ticket discounts (36%)\nGiveaways (9%)\nConcessions (7%)",
        "\"Grab your crew — group deals + collectibles\"",
        "Group ticket landing page; giveaway nights; F&B bundle upsells"
    ],
    [
        "Families w/ Adult Kids",
        "11% of fans",
        "Ticket discounts (38%)\nAffordable food (16%)",
        "\"Quality family time at a price you'll love\"",
        "Family multi-pack ticket bundles with F&B credit; email to 45-65 age cohort"
    ],
    [
        "Older Fans (55+)",
        "49% of fans",
        "Ticket discounts (35%)\nParking credits (10%)",
        "\"Easy parking, great seats, great price\"",
        "Direct mail + email; lead with parking ease; convenience-first messaging"
    ],
    [
        "Passionate Fans (Avidity 5)",
        "40% of fans",
        "Giveaways (7%)\nIntrinsic love (33%)",
        "\"Only for true fans — limited edition exclusives\"",
        "Loyalty/early-access campaigns; exclusive bobblehead drops to drive urgency"
    ],
    [
        "Mid-Avidity Fans (2-3)",
        "12% of fans",
        "Ticket discounts (42-49%)",
        "\"The best deal in baseball — limited time\"",
        "Flash sale emails; urgency-driven CTAs; price anchoring vs. other entertainment"
    ],
]

add_styled_table(doc, headers, rows, col_widths=[3.5, 2.0, 3.5, 4.5, 4.5])

# ============================================================
# PAGE 3 — SUPPORTING EVIDENCE & SATISFACTION PROOF POINTS
# ============================================================

doc.add_page_break()
add_heading_styled(doc, "Supporting Evidence: What 9/10 Fans Love (Use in Creative)", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "These satisfaction data points can serve as proof points in ad copy, landing pages, and email subject lines "
    "to reinforce that the experience delivers on its promise:"
)
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY

# Satisfaction proof points table
sat_headers = ["Experience Area", "% Highly Satisfied", "Messaging Proof Point"]
sat_rows = [
    ["Stadium Cleanliness (Concessions areas)", "80%", "\"Cleanest ballpark experience in Florida\""],
    ["Scoreboard Experience", "70%", "\"Best in-game visuals in baseball\""],
    ["Pregame Content", "66%", "\"Arrive early — the show starts before first pitch\""],
    ["Kids Activities", "64%", "\"Kids love it as much as you do\""],
    ["Food & Beverage Selection", "62%", "\"New food options every time you visit\""],
    ["Music & Atmosphere", "60%", "\"The energy is unmatched\""],
    ["Parking Experience", "Avg 9.26/10", "\"In and out, hassle-free\""],
]

add_styled_table(doc, sat_headers, sat_rows, col_widths=[5.0, 3.0, 7.0])

doc.add_paragraph("")  # spacer

# What fans organically talk about
add_heading_styled(doc, "What Fans Organically Mention (Open-Text Themes)", level=2)
p2 = doc.add_paragraph()
run = p2.add_run(
    "When 9/10 fans write in their own words about their experience, these themes dominate. "
    "Use these in messaging to mirror fan language:"
)
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY

theme_headers = ["Theme", "% Mentioning", "Implication for Messaging"]
theme_rows = [
    ["Stadium quality & renovations", "44%", "Lead with \"new ballpark feel\" in creative"],
    ["Friendly, helpful staff", "23%", "Emphasize welcoming atmosphere in first-timer campaigns"],
    ["Food & beverage experience", "21%", "Feature F&B in social content; supports value messaging"],
    ["Team performance / wins", "13%", "Use win streaks as urgency triggers (\"See them while they're hot\")"],
    ["Great value for money", "9%", "Validate discount messaging with fan sentiment"],
    ["Easy parking & entry", "8%", "Critical for 55+ segment; use as friction-reducer in CTAs"],
]

add_styled_table(doc, theme_headers, theme_rows, col_widths=[4.5, 2.5, 8.0])

doc.add_paragraph("")  # spacer

# ============================================================
# PAGE 4 — SCALING STRATEGY: HOW TO TARGET THESE SEGMENTS
# ============================================================

doc.add_page_break()
add_heading_styled(doc, "Scaling Strategy: How to Target These Segments", level=1)

# Framework intro
fw_intro = doc.add_paragraph()
run = fw_intro.add_run(
    "The VoC data tells us WHAT motivates each fan profile. To grow ticket page visits at scale, "
    "we apply these insights to paid media campaigns that reach far beyond past survey respondents. "
    "The survey is the evidence layer; paid media is the execution layer."
)
run.font.size = Pt(10)
run.font.color.rgb = DARK_GRAY

doc.add_paragraph("")

# --- Strategy 1: Paid Social ---
add_heading_styled(doc, "1. Paid Social (Primary Scale Channel)", level=2)
ps_intro = doc.add_paragraph()
run = ps_intro.add_run(
    "Use Meta (Facebook/Instagram) and TikTok to reach broad audiences matching our high-satisfaction "
    "fan profiles. VoC data informs the creative messaging; platform targeting handles reach."
)
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY

ps_headers = ["Segment", "Platform Targeting", "Creative Angle (Informed by VoC)", "Scale Potential"]
ps_rows = [
    [
        "Women 25-54",
        "Meta: Women, 25-54, Tampa Bay DMA, Interest: MLB/Sports Events",
        "Lead with safety, comfort, clean stadium + ticket deal CTA",
        "HIGH — broad demo, strong platform targeting"
    ],
    [
        "Young Adults 18-34",
        "TikTok + IG Reels: 18-34, Interest: Baseball/Sports\nMeta: 18-34 Lookalike from CRM seed",
        "Giveaway/collectible visuals + \"deals you can't miss\" promo code",
        "HIGH — native platform audience"
    ],
    [
        "Hispanic/Latino",
        "Meta: Hispanic affinity, Tampa DMA, Spanish-language + bilingual\nTikTok: Latino community interest",
        "Heritage night content + exclusive discount offers; bilingual creative",
        "MEDIUM — requires bilingual creative investment"
    ],
    [
        "Black/African American",
        "Meta: Multicultural affinity, Tampa DMA\nEvent-based targeting around theme nights",
        "Theme night event ads (culture, music, community) + parking perk",
        "MEDIUM — event-triggered campaigns"
    ],
    [
        "Friends Groups",
        "Meta: 21-40, Interest: Nightlife/Group Activities/Sports\nIG Stories: shareable group content",
        "\"Grab your crew\" — group deals + giveaway nights + social F&B content",
        "HIGH — social platforms amplify group intent"
    ],
]

add_styled_table(doc, ps_headers, ps_rows, col_widths=[3.0, 5.0, 5.0, 3.0])

doc.add_paragraph("")

# --- Strategy 2: Paid Search/Display ---
add_heading_styled(doc, "2. Paid Search & Display (Capture Active Intent)", level=2)
psd_intro = doc.add_paragraph()
run = psd_intro.add_run(
    "Google Ads and programmatic display capture fans actively searching for tickets or showing "
    "sports/entertainment intent. VoC insights shape ad copy to match what resonates with each profile."
)
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY

psd_headers = ["Segment", "Targeting Method", "Ad Copy Angle (VoC-Informed)", "When to Deploy"]
psd_rows = [
    [
        "Deal Seekers (Mid-Avidity)",
        "Search: \"cheap Rays tickets,\" \"Rays deals,\" \"baseball tickets Tampa\"\nDisplay: In-market Sports Events",
        "Price-first: \"Starting at $10 — Tampa's best entertainment deal\"",
        "Always-on; increase bids during flash sales"
    ],
    [
        "Older Fans 55+",
        "Search: \"Rays game parking,\" \"Rays tickets tonight\"\nDisplay: Age 55+, Tampa DMA, Sports affinity",
        "Convenience-first: \"Free parking + great seats from $XX\"",
        "Weekday games; afternoon starts"
    ],
    [
        "Solo Baseball Fans",
        "Search: \"Rays single tickets,\" \"Rays game tonight\"\nDisplay: In-market MLB, single-ticket modifiers",
        "Product-first: \"Tonight's game. Your seat. From $10.\"",
        "Day-of and day-before game triggers"
    ],
    [
        "Families",
        "Search: \"family things to do Tampa,\" \"kids activities Tampa Bay\"\nDisplay: In-market Family Activities + Sports",
        "Value + kids: \"A night the whole family will love — affordable packages\"",
        "Weekends; school breaks; fireworks nights"
    ],
]

add_styled_table(doc, psd_headers, psd_rows, col_widths=[3.0, 5.5, 5.0, 3.5])

doc.add_paragraph("")

# --- Strategy 3: CRM/Email (Supporting) ---
add_heading_styled(doc, "3. CRM & Email (Retargeting Known Fans)", level=2)
crm_intro = doc.add_paragraph()
run = crm_intro.add_run(
    "CRM campaigns serve as the retargeting layer for fans already in the database. Use segment-specific "
    "messaging and as seed lists for paid social Lookalike audiences to extend reach."
)
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY

crm_headers = ["Segment", "CRM Role", "Message Personalization"]
crm_rows = [
    ["Passionate Fans (Avidity 5)", "Loyalty + early access emails", "\"Exclusive drop — limited edition collectible game\""],
    ["Mid-Avidity (2-3)", "Flash sale + urgency triggers", "\"48-hour deal — lowest price of the season\""],
    ["Past Attendees w/ Adult Kids", "Family multi-pack offers", "\"Bring the family back — F&B credit included\""],
    ["All segments", "Seed lists for Meta/TikTok Lookalike audiences", "Upload CRM segment → 1-2% Lookalike → scale reach"],
]

add_styled_table(doc, crm_headers, crm_rows, col_widths=[4.5, 4.5, 7.0])

# ============================================================
# PAGE 5 — MEASUREMENT & NEXT STEPS
# ============================================================

doc.add_paragraph("")
add_heading_styled(doc, "Measurement Framework", level=2)

meas = doc.add_paragraph()
run = meas.add_run(
    "Every campaign should drive to a UTM-tagged ticket page URL to measure segment-level performance:"
)
run.font.size = Pt(9.5)
run.font.color.rgb = DARK_GRAY

meas_headers = ["Metric", "What It Tells Us", "Target"]
meas_rows = [
    ["Ticket Page Visits (by UTM segment)", "Which segments are responding to messaging", "Establish baseline, grow 15%+ MoM"],
    ["Click-Through Rate (by segment)", "Which message hooks resonate most", "Above platform benchmark (1.5%+ social, 3%+ search)"],
    ["Cost Per Page Visit (by segment)", "Efficiency of spend per audience", "Optimize toward lowest CPV segments"],
    ["Conversion Rate (page visit → purchase)", "Which segments buy after visiting", "Identifies highest-ROI segments for budget increase"],
]

add_styled_table(doc, meas_headers, meas_rows, col_widths=[4.5, 5.5, 5.0])

doc.add_paragraph("")

add_heading_styled(doc, "Recommended Next Steps", level=2)
steps = [
    "Build paid social campaigns for top 3 highest-reach segments (Women 25-54, Young Adults 18-34, Friends Groups) using VoC-informed creative angles",
    "Launch always-on paid search campaigns for deal-seekers and solo fans with day-of game triggers",
    "Upload CRM segment lists to Meta as Custom Audiences; build 1% Lookalikes to scale beyond known fans",
    "Tag all campaign URLs with segment-level UTMs for ticket page visit attribution",
    "Review performance after 4 weeks; shift budget toward highest page-visit-per-dollar segments",
]

for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. {step}")
    run.font.size = Pt(9.5)
    run.font.color.rgb = DARK_GRAY

doc.add_paragraph("")

# Footer note
add_heading_styled(doc, "Data Notes", level=2)
notes = doc.add_paragraph()
run = notes.add_run(
    "Source: TBRDP_DW_DEV.IM_RPT.V_SBL_QUALTRICS_VOC_POST_ATTENDANCE_FULL_CORTEX_AI\n"
    "Filter: Season = 2026, OVERALL_NUMRAT IN (9, 10)  |  n = 8,332 respondents\n"
    "Incentive ranking data available for 1,195 respondents (14.3%); open-text data for 6,430 (77%)\n"
    "Visit frequency fields were unpopulated for 2026; Team Avidity used as engagement proxy.\n\n"
    "Strategic Note: VoC survey data serves as the research/evidence layer justifying messaging decisions. "
    "Campaign execution scales beyond survey respondents via paid media audience targeting in Tampa Bay DMA."
)
run.font.size = Pt(8)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True

# --- Save ---
doc.save(OUT)
print(f"Document saved to: {OUT}")
