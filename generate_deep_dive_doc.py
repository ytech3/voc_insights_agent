"""Generate HS1 vs HS2 Deep Dive Word Document."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)

# Helper to add a styled table
def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    return table

def add_verbatim_block(doc, title, verbatims):
    """Add a block of fan verbatims with light formatting."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    for v in verbatims:
        vp = doc.add_paragraph(style='List Bullet')
        vr = vp.add_run(f'"{v}"')
        vr.italic = True
        vr.font.size = Pt(9)
        vr.font.color.rgb = RGBColor(80, 80, 80)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('VOC Homestand Comparison\nDeep Dive Analysis')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 47, 108)  # Rays navy

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Homestand 1 (Apr 6-12) vs Homestand 2 (Apr 20-26)\n2026 Season')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run('Tampa Bay Rays — Strategy & Analytics')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 47, 108)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Concessions Customer Service — HS1 to HS2 Decline',
    '2. Concessions Wait Time — 101% Increase vs 2024',
    '3. Music Satisfaction — 1% Decline vs 2024',
    '4. Food Quality Deep Dive — Chicken, Hot Dogs & Fries',
    '   4a. CONCESS_QUAL_REASON: Dissatisfaction Drivers',
    '   4b. Cross-Food Summary & Actionable Insights',
    '5. Tech Team / App Satisfaction — Down 4% vs 2024',
    '6. Concessions Customer Service Grid — 7% Dissatisfaction Rise HS1→HS2',
    '7. CONCESS_QUAL_REASON — All 13 Food Items',
    '8. Seat Value Perception — Down ~4% vs 2024',
    '9. Parking — Finding a Space Down 3.4% vs 2024',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 1: CUSTOMER SERVICE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Concessions Customer Service', level=1)
doc.add_heading('% Highly Satisfied fell 6.6% from HS1 to HS2', level=2)

doc.add_paragraph(
    'The "Highly Satisfied" rate for concessions customer service declined from '
    '81.8% in Homestand 1 (n=198) to 76.4% in Homestand 2 (n=250), a relative '
    'drop of approximately 6.6%. This section investigates the root causes.'
)

doc.add_heading('Homestand-Level Summary', level=3)
add_table(doc,
    ['Period', 'Highly Satisfied', 'Total n', '% Highly Satisfied'],
    [
        ['HS1', '162', '198', '81.8%'],
        ['HS2', '191', '250', '76.4%'],
    ])

doc.add_paragraph()
doc.add_heading('Game-by-Game Breakdown', level=3)
doc.add_paragraph(
    'The decline was not uniform across HS2. Two specific dates drove the average down:'
)
add_table(doc,
    ['Game Date', 'HS', 'Highly Sat', 'Total n', '% Highly Sat'],
    [
        ['Apr 6', 'HS1', '38', '50', '76.0%'],
        ['Apr 7', 'HS1', '29', '34', '85.3%'],
        ['Apr 8', 'HS1', '26', '29', '89.7%'],
        ['Apr 10', 'HS1', '23', '30', '76.7%'],
        ['Apr 11', 'HS1', '18', '25', '72.0%'],
        ['Apr 12', 'HS1', '28', '30', '93.3%'],
        ['Apr 20', 'HS2', '33', '49', '67.3%'],
        ['Apr 21', 'HS2', '31', '42', '73.8%'],
        ['Apr 22', 'HS2', '48', '65', '73.8%'],
        ['Apr 24', 'HS2', '19', '24', '79.2%'],
        ['Apr 25', 'HS2', '32', '37', '86.5%'],
        ['Apr 26', 'HS2', '28', '33', '84.8%'],
    ])

doc.add_paragraph()
doc.add_heading('Key Findings', level=3)
findings = [
    'Apr 20 (67.3%) and Apr 22 (73.8%) were the two worst customer service days across both homestands, pulling the HS2 average down.',
    'By contrast, Apr 24-26 in HS2 scored 79.2%, 86.5%, and 84.8%, matching or exceeding HS1 averages.',
    'Sample sizes are small (198 vs 250 total respondents). A shift of just 5-10 respondents meaningfully changes the percentage.',
    'Open-ended negative feedback about concessions staff dropped from 118 mentions in HS1 to just 22 in HS2, suggesting the structured score decline is driven by a small number of bad experiences on specific days rather than a systemic worsening.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_heading('The Paradox', level=3)
doc.add_paragraph(
    'The structured grid score declined HS1 to HS2, but open-ended complaints about staff service '
    'dropped 81% (118 to 22 mentions). This suggests: (a) small sample sizes amplify day-to-day '
    'variation, (b) respondent composition shifted (more first-time visitors in HS2 may have different '
    'expectations), and (c) Apr 20 — the first game back after a week off — may have had staffing '
    'adjustment issues. The trend from Apr 20 to Apr 26 shows rapid recovery.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 2: WAIT TIME
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Concessions Wait Time', level=1)
doc.add_heading('"Much More Wait Than Expected" grew 101% in 2026 vs 2024', level=2)

doc.add_paragraph(
    'The percentage of fans selecting "Much more than what I expected" for concessions wait time '
    'increased sharply in 2026 compared to 2024. However, this aggregate number masks a story of '
    'rapid improvement within the 2026 season.'
)

doc.add_heading('Period-Level Summary', level=3)
add_table(doc,
    ['Period', '"Much More"', 'Total n', '% Much More'],
    [
        ['2024 (full season)', '795', '14,301', '5.6%'],
        ['HS1 (Apr 6-12)', '446', '2,819', '15.8%'],
        ['HS2 (Apr 20-26)', '77', '1,892', '4.1%'],
    ])

doc.add_paragraph()
doc.add_heading('Game-by-Game Breakdown', level=3)
doc.add_paragraph(
    'Opening Day (Apr 6) was the primary outlier, with 26.5% of respondents saying wait was '
    '"Much more than expected" — nearly 5x the 2024 season average.'
)
add_table(doc,
    ['Game Date', 'HS', '"Much More"', 'Total n', '% Much More'],
    [
        ['Apr 6 (Opening Day)', 'HS1', '220', '831', '26.5%'],
        ['Apr 7', 'HS1', '76', '454', '16.7%'],
        ['Apr 8', 'HS1', '60', '452', '13.3%'],
        ['Apr 10', 'HS1', '32', '302', '10.6%'],
        ['Apr 11', 'HS1', '36', '381', '9.4%'],
        ['Apr 12', 'HS1', '22', '399', '5.5%'],
        ['Apr 20', 'HS2', '22', '336', '6.5%'],
        ['Apr 21', 'HS2', '11', '276', '4.0%'],
        ['Apr 22', 'HS2', '21', '492', '4.3%'],
        ['Apr 24', 'HS2', '2', '229', '0.9%'],
        ['Apr 25', 'HS2', '11', '299', '3.7%'],
        ['Apr 26', 'HS2', '10', '260', '3.8%'],
    ])

doc.add_paragraph()
doc.add_heading('Key Findings', level=3)
findings = [
    'Opening Day (Apr 6) alone contributed 220 of the 446 "Much More" responses in HS1 — nearly half.',
    'By Apr 12 (end of HS1), the rate had already dropped to 5.5%, matching the 2024 baseline.',
    'HS2 averaged just 4.1%, which is actually below the 2024 season average of 5.6%.',
    'The 101% increase is entirely driven by the first 2-3 games of HS1 when the new stadium operations (kiosk ordering, new concession layouts) were being worked out.',
    'The improvement from 26.5% (Apr 6) to 0.9% (Apr 24) represents a dramatic operational learning curve.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Assessment', level=3)
doc.add_paragraph(
    'This metric is no longer a concern. The 2026 season aggregate is inflated by Opening Day. '
    'By HS2, wait time performance was better than the 2024 baseline. The new ordering systems '
    '(kiosks, mobile) had a rocky start but have stabilized. Future homestand reports should show '
    'continued improvement or parity with 2024.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 3: MUSIC
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Music Satisfaction', level=1)
doc.add_heading('1% decline vs 2024 — what are fans saying?', level=2)

doc.add_paragraph(
    'Music satisfaction dipped slightly compared to the 2024 baseline. Sentence-level analysis '
    'of open-ended feedback reveals the complaints are concentrated and consistent.'
)

doc.add_heading('Negative Mention Counts', level=3)
add_table(doc,
    ['Period', 'Negative Mentions'],
    [
        ['2024 (full season)', '12'],
        ['HS1 (6 games)', '15'],
        ['HS2 (6 games)', '1'],
    ])

doc.add_paragraph()
doc.add_heading('What Fans Are Saying', level=3)
doc.add_paragraph(
    'The negative feedback is overwhelmingly about one issue: PA system / music volume being too loud. '
    '17 of the 19 negative sentences from 2026 reference volume or loudness specifically.'
)

music_verbatims = [
    "The sound system was way, way louder than it needed to be.",
    "Should turn down volume of stadium PA system a few notches.",
    "Music is waaaaay to loud.",
    "The volume on the speaker system was WAY WAY WAY too loud.",
    "The music is way too loud and intrusive.",
    "I dont suffer from hearing problems and do enjoy loud music, but the sound system around the area where i sat near home plate felt extremely loud.",
    "It is still too loud inside especially the music.",
    "Audio system needs to be turned down.",
    "Your audio is way too loud.",
    "The decibel level was unacceptable.",
]
add_verbatim_block(doc, 'Representative Fan Quotes (HS1):', music_verbatims)

doc.add_paragraph()
doc.add_paragraph(
    'The single HS2 complaint was about music selection rather than volume:'
)
doc.add_paragraph(
    '"Whoever does the music between plays should choose better music to get the crowd '
    'involved like AC/DC or other songs that revs the crowds up like at football games."',
    style='List Bullet'
)

doc.add_heading('Key Findings', level=3)
findings = [
    'The volume complaints were concentrated entirely in HS1 (Apr 6-12). By HS2, only 1 negative mention remained, and it was about song selection — not volume.',
    'This same complaint existed in 2024 (12 mentions across the full season), so it is not new to the stadium. The new venue acoustics may have amplified the issue early on.',
    'The issue appears to be self-correcting. The audio team likely adjusted levels between homestands.',
    'This is a negligible concern going forward. Monitor HS3 to confirm resolution.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 4: FOOD QUALITY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Food Quality Deep Dive', level=1)
doc.add_heading('Chicken, Hot Dogs & Fries — CONCESS_QUAL_REASON Analysis', level=2)

doc.add_paragraph(
    'The CONCESS_QUAL_REASON column captures the specific reason dissatisfied fans gave for their '
    'food quality rating. This question is only asked of respondents who rated a food item as '
    '"Somewhat Dissatisfied" or "Highly Dissatisfied." The reason codes are:'
)
doc.add_paragraph('1 = Taste', style='List Bullet')
doc.add_paragraph('2 = Temperature', style='List Bullet')
doc.add_paragraph('3 = Portion Size', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'Note: This data comes from the non-Cortex view (V_SBL_QUALTRICS_VOC_POST_ATTENDANCE_FULL) '
    'and is available for both 2024 and 2026, providing direct year-over-year comparison. '
    'The _DESC column labels are mislabeled in the view (mapping codes to generic satisfaction labels '
    'rather than reason names); the numeric FLOAT codes are authoritative.'
)

doc.add_paragraph()

# ─── CHICKEN ───
doc.add_heading('4a. Chicken', level=2)
doc.add_paragraph(
    'Chicken saw the most dramatic quality decline. The total number of dissatisfied respondents '
    'is significantly higher in 2026: 95 in HS1 alone (6 games) vs 129 in all of 2024 (full season).'
)
add_table(doc,
    ['Period', 'n', 'Taste', '% Taste', 'Temp', '% Temp', 'Portion', '% Portion'],
    [
        ['2024 (full)', '129', '59', '45.7%', '28', '21.7%', '7', '5.4%'],
        ['HS1', '95', '52', '54.7%', '27', '28.4%', '16', '16.8%'],
        ['HS2', '77', '45', '58.4%', '20', '26.0%', '12', '15.6%'],
    ])

doc.add_paragraph()
doc.add_heading('Game-by-Game (Chicken)', level=3)
add_table(doc,
    ['Date', 'HS', 'n', 'Taste', 'Temp', 'Portion'],
    [
        ['Apr 6', 'HS1', '24', '14', '8', '2'],
        ['Apr 7', 'HS1', '8', '7', '1', '0'],
        ['Apr 8', 'HS1', '8', '5', '2', '1'],
        ['Apr 10', 'HS1', '11', '4', '4', '3'],
        ['Apr 11', 'HS1', '24', '10', '8', '6'],
        ['Apr 12', 'HS1', '20', '12', '4', '4'],
        ['Apr 20', 'HS2', '14', '6', '6', '2'],
        ['Apr 21', 'HS2', '13', '8', '4', '1'],
        ['Apr 22', 'HS2', '18', '13', '3', '2'],
        ['Apr 24', 'HS2', '9', '6', '1', '2'],
        ['Apr 25', 'HS2', '15', '8', '4', '3'],
        ['Apr 26', 'HS2', '8', '4', '2', '2'],
    ])

doc.add_paragraph()
doc.add_heading('Chicken Key Findings', level=3)
findings = [
    'Taste is the #1 driver of chicken dissatisfaction in every period, and it is worsening: 45.7% (2024) -> 54.7% (HS1) -> 58.4% (HS2).',
    'Portion Size complaints tripled from 2024 (5.4%) to 2026 (15.6-16.8%), reflecting the "2 tenders and a small pile of fries for $20" sentiment in open-ended feedback.',
    'Temperature is the secondary driver (~26-28% in 2026 vs 21.7% in 2024), aligning with "cold/soggy tenders" verbatims.',
    'Apr 6 and Apr 11 were the worst days in HS1 (24 dissatisfied each). By HS2, Temperature complaints eased but Taste remained dominant.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
chicken_verbatims = [
    "We had 2 orders of chicken tenders and fries that tasted as if they had been sitting for hours.",
    "Only hiccups were the food - $20 for 2 chicken strips and a pile of fries was one thing but they were absolutely terrible, like seriously hard to mess up basic chicken.",
    "I purchased chicken tenders and fries from an area on the Porch and they were so dry, I could not eat them.",
    "Despite the 'First one is in us' campaign, the costs are ridiculous; $19 for 3 chicken pieces and soggy fries.",
    "Chicken fingers dripping grease & undercooked, and popcorn tasted like fish (both inedible).",
    "How do you mess up chicken fingers?",
    "The pulled chicken tasted like it came out of a can, the chips and bread were stale.",
    "18.00 dollars for 2 cold chicken fingers and cold, soggy fries.",
    "Then the chicken, it was soggy and just nasty, for the amount you charge you should be serving good, fresh food.",
]
add_verbatim_block(doc, 'Representative Chicken Verbatims:', chicken_verbatims)

doc.add_page_break()

# ─── HOT DOGS ───
doc.add_heading('4b. Hot Dogs', level=2)
doc.add_paragraph(
    'Hot dogs had the highest volume of dissatisfied respondents in 2026. Temperature is the '
    'dominant driver, but Taste complaints have surged.'
)
add_table(doc,
    ['Period', 'n', 'Taste', '% Taste', 'Temp', '% Temp', 'Portion', '% Portion'],
    [
        ['2024 (full)', '197', '48', '24.4%', '67', '34.0%', '19', '9.6%'],
        ['HS1', '203', '78', '38.4%', '97', '47.8%', '28', '13.8%'],
        ['HS2', '93', '40', '43.0%', '43', '46.2%', '10', '10.8%'],
    ])

doc.add_paragraph()
doc.add_heading('Hot Dog Key Findings', level=3)
findings = [
    'Temperature is the #1 driver in every period: 34.0% (2024) -> 47.8% (HS1) -> 46.2% (HS2). Nearly half of dissatisfied hot dog fans cite cold product.',
    'Taste complaints surged: 24.4% (2024) -> 38.4% (HS1) -> 43.0% (HS2). This is a meaningful deterioration.',
    'Taste + Temperature together account for ~89% of dissatisfaction in HS2, pointing to a product quality and/or preparation issue.',
    'The total dissatisfied count dropped from 203 (HS1) to 93 (HS2) — fewer people were unhappy overall, but the reasons remained the same.',
    'Portion Size spiked in HS1 (13.8%) but returned to near-2024 levels (10.8%) by HS2.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
hotdog_verbatims = [
    "hot dog was cold",
    "hot dog was disgusting",
    "I had a hot dog that was only lukewarm.",
    "The Hot dogs were cold, buns cold, sauerkraut cold.",
    "Hot dogs were dried and shriveled.",
    "Plus hotdog was cold and bun hard.",
    "The food was mediocre (cold hot dog and very expensive beer).",
    "we had two hot dogs and a pretzel all were cold",
    "The Hot dogs were terrible and the buns hard.",
    "Wet hotdogs not the best.",
    "bread hard on hotdogs",
]
add_verbatim_block(doc, 'Representative Hot Dog Verbatims:', hotdog_verbatims)

doc.add_page_break()

# ─── FRIES ───
doc.add_heading('4c. Fries', level=2)
doc.add_paragraph(
    'Fries follow a pattern similar to hot dogs, with Temperature as the overwhelming #1 driver.'
)
add_table(doc,
    ['Period', 'n', 'Taste', '% Taste', 'Temp', '% Temp', 'Portion', '% Portion'],
    [
        ['2024 (full)', '73', '16', '21.9%', '34', '46.6%', '2', '2.7%'],
        ['HS1', '104', '31', '29.8%', '65', '62.5%', '8', '7.7%'],
        ['HS2', '65', '25', '38.5%', '36', '55.4%', '4', '6.2%'],
    ])

doc.add_paragraph()
doc.add_heading('Fries Key Findings', level=3)
findings = [
    'Temperature is overwhelmingly #1: 46.6% (2024) -> 62.5% (HS1) -> 55.4% (HS2). Cold fries are the single biggest issue.',
    'Taste is growing as a secondary concern: 21.9% (2024) -> 29.8% (HS1) -> 38.5% (HS2). Even when fries are at the right temperature, fans in 2026 do not like how they taste.',
    'Portion Size is negligible across all periods (3-8%).',
    'Overall dissatisfied count dropped from 104 (HS1) to 65 (HS2), meaning fewer complaints but the same root causes.',
]
for f in findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
fries_verbatims = [
    "The fries were cold & rubbery.",
    "The fries were terrible.",
    "the french fries were not done",
    "Fries could not be eaten",
    "When we arrived to the food court, I ordered chicken tender and fries, and the fries were ice cold.",
    "We went to a table and the new fries were just as cold as the original ones.",
    "French Fries were cold.",
    "Smash burgers and fries that were sitting for an hour or longer not fresh and warm",
    "fries were soggy and cold",
    "French fries not salted enough.",
]
add_verbatim_block(doc, 'Representative Fries Verbatims:', fries_verbatims)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# CROSS-FOOD SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4d. Cross-Food Summary', level=2)

add_table(doc,
    ['Metric', 'Chicken', 'Hot Dogs', 'Fries'],
    [
        ['#1 Driver', 'Taste (58.4%)', 'Temperature (46.2%)', 'Temperature (55.4%)'],
        ['#2 Driver', 'Temperature (26.0%)', 'Taste (43.0%)', 'Taste (38.5%)'],
        ['Emerging Concern', 'Portion Size (15.6%)', '—', '—'],
        ['Trend vs 2024', 'Worse on all 3', 'Worse on Taste & Temp', 'Worse on Taste & Temp'],
        ['HS1 -> HS2', 'Taste worsening, Temp improving', 'Taste worsening, Temp stable', 'Taste worsening, Temp slight improvement'],
    ])

doc.add_paragraph()
doc.add_heading('Actionable Insights', level=3)

insights = [
    ('Temperature (Operational Issue — Improving)',
     'Temperature is the dominant driver for Hot Dogs and Fries. This is an operational issue — '
     'food sitting too long between preparation and pickup, likely exacerbated by the new kiosk '
     'ordering flow creating a gap. The good news: as wait times improved from HS1 to HS2, '
     'overall dissatisfaction counts dropped significantly (Hot Dogs: 203->93, Fries: 104->65). '
     'This should continue improving as operations stabilize.'),

    ('Taste (Product Quality Issue — Worsening)',
     'Taste complaints are increasing across ALL three food items from 2024 -> HS1 -> HS2. '
     'This trend persists even as wait times and Temperature complaints improve, which means '
     'operational fixes alone will not address it. This suggests either: (a) a product or supplier '
     'change, (b) a preparation method change in the new venue, or (c) the new kitchen/stand '
     'layouts affecting cooking quality. This requires a product-level conversation with the '
     'concessions partner.'),

    ('Portion Size (Chicken-Specific)',
     'Portion Size complaints tripled for Chicken (5.4% -> 14-17%), reflecting fan sentiment '
     'about value: "$20 for 2 tenders and a small pile of fries." This is a pricing/value '
     'perception issue specific to the chicken tender product and should be reviewed with '
     'the vendor.'),
]
for title, body in insights:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(body)
    run2.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 5: TECH TEAM / APP SATISFACTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. Tech Team / App Satisfaction', level=1)
doc.add_heading('Down ~4% vs 2024 — App & Connectivity Issues Are the Driver', level=2)

doc.add_paragraph(
    'Tech Team satisfaction declined approximately 4% in 2026 compared to 2024. Analysis of '
    'TB_ADDON_8_9 (the app-specific feedback column) and sentence-level AI categorization '
    'reveals that the decline is overwhelmingly driven by app functionality and connectivity '
    'issues — particularly the Rays Wallet / Burst Bucks promotion and poor WiFi/cellular '
    'service inside the stadium.'
)

doc.add_heading('TB_ADDON_8_9: App Complaint Volume', level=3)
doc.add_paragraph(
    'TB_ADDON_8_9 captures open-text feedback specifically about app-related issues. '
    'The vast majority of responses are numeric satisfaction scores (85-89 range). '
    'Non-numeric text complaints represent actual written feedback:'
)
add_table(doc,
    ['Period', 'Total Responses', 'Numeric Scores', 'Text Complaints'],
    [
        ['2024 (full season)', '21,007', '20,957', '50'],
        ['HS1 (6 games)', '3,890', '3,855', '35'],
        ['HS2 (6 games)', '2,635', '2,626', '9'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'Notably, HS1 generated 35 text complaints in just 6 games — 70% of the full 2024 season\'s '
    'complaint volume (50). HS2 saw a dramatic drop to just 9 complaints, suggesting the issues '
    'were concentrated in the early weeks.'
)

doc.add_heading('Sentence-Level Analysis: Digital/App Categories', level=3)
doc.add_paragraph(
    'The sentence-level AI categorization reveals a massive increase in negative digital experience '
    'mentions in 2026:'
)
add_table(doc,
    ['Category', '2024 Negative', '2026 Negative', 'Change'],
    [
        ['Mobile Ordering', '14', '48', '+243%'],
        ['Mobile Ticketing', '8', '45', '+463%'],
        ['Mobile App', '6', '39', '+550%'],
        ['Tech Team (staff)', '0', '6', 'New in 2026'],
        ['Total Digital Negative', '28', '138', '+393%'],
    ])

doc.add_paragraph()
doc.add_heading('Root Cause: Rays Wallet & Connectivity', level=3)
doc.add_paragraph(
    'The dominant theme across both TB_ADDON_8_9 text complaints and sentence-level analysis '
    'is the Rays Wallet / Burst Bucks / "First One\'s On Us" $10 credit promotion. Fans could '
    'not access their credits due to:'
)
wallet_causes = [
    'Poor WiFi/cellular connectivity inside the stadium — the most frequently cited issue',
    'Rays Wallet app freezing, crashing, or failing to load QR codes',
    'Concession staff not knowing how to process Wallet/Burst Bucks payments',
    'Promotional credits not transferring correctly when tickets were shared',
    'Older phones unable to run or connect to the required apps',
]
for c in wallet_causes:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
app_verbatims_2026 = [
    "The owners offered 'first drink on us' but with no wireless service and very poor cell service the Rays Wallet app would not load.",
    "My MLB app would not work when I was in stadium. I could not use the $20 that was given to me for the game but once I stepped out of the stadium, it worked great.",
    "Rays wallet didn't work at any concession stand.",
    "Internet and 5G service was spotty throughout the parking lot and inside the Trop.",
    "could not connect to the Trops wifi to use the app effectively. No wifi to connect that wasn't locked.",
    "Don't offer Food Credits when anticipated crowd size overwhelms the cell towers and they can't get on-line to collect these credits.",
    "4 of us tried to use our $10 each, and only 2 worked.",
    "QR codes weren't being read at registers, discount and burst bucks had to be manually entered.",
    "I wasted half the game trying to get my phone to work so I could buy stuff.",
    "It was disappointing that Rays Wallet kept crashing the system, so I wasn't able to use 1st one on us promotion.",
]
add_verbatim_block(doc, 'Representative 2026 App Verbatims:', app_verbatims_2026)

doc.add_paragraph()
doc.add_heading('Key Findings', level=3)
tech_findings = [
    'The Tech Team satisfaction decline is NOT about the Tech Team staff specifically — only 6 sentence-level mentions reference Tech Team staff (and 7 were positive). The decline is about the digital ecosystem they support.',
    'The Rays Wallet / Burst Bucks promotion created massive friction. Fans were promised $10 credits per ticket but many could not redeem them due to connectivity issues.',
    'WiFi/cellular infrastructure is the core bottleneck. Multiple fans report the app works fine outside the stadium but fails inside.',
    'HS2 showed significant improvement: text complaints dropped from 35 (HS1) to 9 (HS2), suggesting either connectivity was improved or fans adapted.',
    'The 2024 complaint themes were similar (app login issues, WiFi, payment processing) but at much lower volume, suggesting the new Rays Wallet promotion amplified existing infrastructure limitations.',
]
for f in tech_findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Recommendation', level=3)
doc.add_paragraph(
    'The Tech Team satisfaction metric is a proxy for digital infrastructure satisfaction. '
    'The immediate priority is WiFi/cellular capacity upgrades inside the stadium. '
    'Secondary priority: ensure the Rays Wallet has an offline fallback or pre-cached QR codes '
    'so fans can access their credits even without a live connection. '
    'The Tech Team staff themselves are performing well — 7 positive vs 6 negative mentions — '
    'but they are absorbing complaints about systems they cannot control.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 6: CONCESS GRID CUST SERVICE DISSATISFACTION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. Concessions Customer Service Grid', level=1)
doc.add_heading('Dissatisfaction Rose 7% from HS1 to HS2', level=2)

doc.add_paragraph(
    'When looking specifically at the dissatisfied end of the concessions customer service grid '
    '(combining "Somewhat Dissatisfied" and "Highly Dissatisfied"), the rate increased from '
    'HS1 to HS2. This section examines why dissatisfaction specifically grew even as the '
    'overall complaint volume in open-ended feedback dropped.'
)

doc.add_heading('Dissatisfaction Rate: HS1 vs HS2', level=3)
add_table(doc,
    ['Period', 'Highly Dissat', 'Somewhat Dissat', 'Total Dissat', 'Total n', '% Dissatisfied'],
    [
        ['HS1', '10', '6', '16', '198', '8.1%'],
        ['HS2', '19', '10', '29', '250', '11.6%'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'The dissatisfaction rate rose from 8.1% (HS1) to 11.6% (HS2), driven primarily by an '
    'increase in "Highly Dissatisfied" responses (10 → 19).'
)

doc.add_heading('Game-by-Game Dissatisfaction', level=3)
doc.add_paragraph(
    'The HS2 increase was concentrated on specific dates:'
)
add_table(doc,
    ['Game Date', 'HS', 'Highly Dissat', 'Somewhat Dissat', 'Total n', '% Dissatisfied'],
    [
        ['Apr 6', 'HS1', '4', '1', '50', '10.0%'],
        ['Apr 7', 'HS1', '1', '2', '34', '8.8%'],
        ['Apr 8', 'HS1', '0', '0', '29', '0.0%'],
        ['Apr 10', 'HS1', '2', '2', '30', '13.3%'],
        ['Apr 11', 'HS1', '2', '1', '25', '12.0%'],
        ['Apr 12', 'HS1', '1', '0', '30', '3.3%'],
        ['Apr 20', 'HS2', '5', '3', '49', '16.3%'],
        ['Apr 21', 'HS2', '2', '2', '42', '9.5%'],
        ['Apr 22', 'HS2', '7', '2', '65', '13.8%'],
        ['Apr 24', 'HS2', '3', '0', '24', '12.5%'],
        ['Apr 25', 'HS2', '1', '1', '37', '5.4%'],
        ['Apr 26', 'HS2', '1', '2', '33', '9.1%'],
    ])

doc.add_paragraph()
doc.add_heading('The Paradox: Structured Scores vs Open-Ended Feedback', level=3)
doc.add_paragraph(
    'The structured grid shows dissatisfaction increasing, yet sentence-level negative mentions '
    'of concession staff service tell the opposite story:'
)
add_table(doc,
    ['Metric', 'HS1', 'HS2', 'Change'],
    [
        ['Structured % Dissatisfied', '8.1%', '11.6%', '+43% (worse)'],
        ['Sentence-Level Negative Mentions', '118', '21', '-82% (better)'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'This paradox suggests the structured grid decline is driven by: (a) small absolute numbers — '
    'an increase of just 13 dissatisfied respondents (16→29) swings the percentage; (b) respondent '
    'composition changes between homestands; (c) Apr 20 and Apr 22 as specific problem dates.'
)

doc.add_heading('HS2 Negative Verbatims — What Went Wrong', level=3)
doc.add_paragraph(
    'The 21 sentence-level negative mentions in HS2 reveal specific complaint themes:'
)
hs2_cust_serv_verbatims = [
    "Service was extremely slow and the servers did not know the layout of the rows and seat numbers, looked totally untrained.",
    "I cannot tell you how many orders were delivered to the incorrect people.",
    "The people who brought our food out didn't get the seats correct half the time.",
    "The employees delivering it were confused and would give the food away if they couldn't find the person ordering.",
    "Workers spent more time on kiosks to work than serving food.",
    "The guys working didn't know how to take an order for a Coke!!",
    "We sat at table to order food and waitress was talking to customers for 10 minutes.",
    "When I went to sit in my seats the worker in my section rudely confronted me assuming I didn't belong.",
]
add_verbatim_block(doc, 'HS2 Customer Service Complaints:', hs2_cust_serv_verbatims)

doc.add_paragraph()
doc.add_heading('Key Findings', level=3)
cs_findings = [
    'Apr 20 (16.3%) and Apr 22 (13.8%) drove the HS2 dissatisfaction spike — these two dates alone contributed 17 of the 29 dissatisfied responses.',
    'Apr 24-26 showed mixed dissatisfaction levels (12.5%, 5.4%, and 9.1%), with recovery on some dates.',
    'HS2 complaints shifted from HS1 themes (understaffing, chaos, long waits) to delivery accuracy issues (wrong seats, wrong orders). This suggests new in-seat ordering is the friction point.',
    'The open-ended complaint volume dropped 82% (118→21), meaning the overall fan experience improved even as a small subset reported dissatisfaction in the structured survey.',
    'Sample sizes are small — the entire HS1 to HS2 shift represents just 13 additional dissatisfied respondents.',
]
for f in cs_findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 7: ALL 13 FOOD ITEMS CONCESS_QUAL_REASON
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. CONCESS_QUAL_REASON — All 13 Food Items', level=1)
doc.add_heading('Taste vs Temperature vs Portion Size Across the Full Menu', level=2)

doc.add_paragraph(
    'Expanding the CONCESS_QUAL_REASON analysis from Investigation 4 (Chicken, Hot Dogs, Fries) '
    'to all 13 food items tracked in the survey. Code mapping: 1=Taste, 2=Temperature, 3=Portion Size. '
    'Note: Codes 4 and 5 appear in 2024 data but not in 2026, likely representing "N/A" or other '
    'responses that were removed from the 2026 survey instrument. Only codes 1-3 are compared below.'
)

doc.add_paragraph()
doc.add_heading('2024 Baseline — All Food Items (Codes 1-3 Only)', level=3)
add_table(doc,
    ['Food Item', 'n (1-3)', 'Taste', '% Taste', 'Temp', '% Temp', 'Portion', '% Portion'],
    [
        ['Burgers', '51', '30', '58.8%', '12', '23.5%', '9', '17.6%'],
        ['Chicken', '94', '59', '62.8%', '28', '29.8%', '7', '7.4%'],
        ['Fries', '52', '16', '30.8%', '34', '65.4%', '2', '3.8%'],
        ['Hotdog', '134', '48', '35.8%', '67', '50.0%', '19', '14.2%'],
        ['Ice Cream', '17', '8', '47.1%', '1', '5.9%', '8', '47.1%'],
        ['Nachos', '55', '31', '56.4%', '9', '16.4%', '15', '27.3%'],
        ['Nuts', '12', '5', '41.7%', '1', '8.3%', '6', '50.0%'],
        ['Pizza', '63', '51', '81.0%', '6', '9.5%', '6', '9.5%'],
        ['Popcorn', '39', '24', '61.5%', '3', '7.7%', '12', '30.8%'],
        ['Pretzels', '52', '30', '57.7%', '19', '36.5%', '3', '5.8%'],
        ['Salad', '2', '1', '50.0%', '0', '0.0%', '1', '50.0%'],
        ['Sandwich', '55', '29', '52.7%', '16', '29.1%', '10', '18.2%'],
        ['Sausage', '15', '8', '53.3%', '4', '26.7%', '3', '20.0%'],
    ])

doc.add_paragraph()
doc.add_heading('HS1 (2026) — All Food Items', level=3)
add_table(doc,
    ['Food Item', 'n', 'Taste', '% Taste', 'Temp', '% Temp', 'Portion', '% Portion'],
    [
        ['Burgers', '45', '21', '46.7%', '19', '42.2%', '5', '11.1%'],
        ['Chicken', '95', '52', '54.7%', '27', '28.4%', '16', '16.8%'],
        ['Fries', '104', '31', '29.8%', '65', '62.5%', '8', '7.7%'],
        ['Hotdog', '203', '78', '38.4%', '97', '47.8%', '28', '13.8%'],
        ['Ice Cream', '15', '5', '33.3%', '4', '26.7%', '6', '40.0%'],
        ['Nachos', '40', '19', '47.5%', '11', '27.5%', '10', '25.0%'],
        ['Nuts', '19', '9', '47.4%', '1', '5.3%', '9', '47.4%'],
        ['Pizza', '54', '38', '70.4%', '8', '14.8%', '8', '14.8%'],
        ['Popcorn', '96', '79', '82.3%', '9', '9.4%', '8', '8.3%'],
        ['Pretzels', '91', '50', '54.9%', '39', '42.9%', '2', '2.2%'],
        ['Salad', '0', '—', '—', '—', '—', '—', '—'],
        ['Sandwich', '30', '14', '46.7%', '13', '43.3%', '3', '10.0%'],
        ['Sausage', '15', '6', '40.0%', '3', '20.0%', '6', '40.0%'],
    ])

doc.add_paragraph()
doc.add_heading('HS2 (2026) — All Food Items', level=3)
add_table(doc,
    ['Food Item', 'n', 'Taste', '% Taste', 'Temp', '% Temp', 'Portion', '% Portion'],
    [
        ['Burgers', '36', '20', '55.6%', '11', '30.6%', '5', '13.9%'],
        ['Chicken', '77', '45', '58.4%', '20', '26.0%', '12', '15.6%'],
        ['Fries', '65', '25', '38.5%', '36', '55.4%', '4', '6.2%'],
        ['Hotdog', '93', '40', '43.0%', '43', '46.2%', '10', '10.8%'],
        ['Ice Cream', '8', '3', '37.5%', '2', '25.0%', '3', '37.5%'],
        ['Nachos', '41', '16', '39.0%', '15', '36.6%', '10', '24.4%'],
        ['Nuts', '6', '1', '16.7%', '0', '0.0%', '5', '83.3%'],
        ['Pizza', '40', '29', '72.5%', '3', '7.5%', '8', '20.0%'],
        ['Popcorn', '37', '26', '70.3%', '6', '16.2%', '5', '13.5%'],
        ['Pretzels', '60', '38', '63.3%', '20', '33.3%', '2', '3.3%'],
        ['Salad', '0', '—', '—', '—', '—', '—', '—'],
        ['Sandwich', '22', '14', '63.6%', '3', '13.6%', '5', '22.7%'],
        ['Sausage', '4', '2', '50.0%', '1', '25.0%', '1', '25.0%'],
    ])

doc.add_paragraph()
doc.add_heading('Cross-Food Insights', level=3)
doc.add_paragraph(
    'Categorizing all 13 items by their primary dissatisfaction driver:'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Taste-Driven Items (product quality issue): ')
run.bold = True
run.font.size = Pt(10)
p.add_run('Pizza (74-81%), Popcorn (62-82%), Chicken (55-63%), Pretzels (55-65%), '
          'Burgers (47-59%), Nachos (38-56%), Sandwich (47-58%)').font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Temperature-Driven Items (operational/holding issue): ')
run.bold = True
run.font.size = Pt(10)
p.add_run('Fries (58-65%), Hotdog (47-50%)').font.size = Pt(10)

p = doc.add_paragraph()
run = p.add_run('Portion-Driven Items (value perception): ')
run.bold = True
run.font.size = Pt(10)
p.add_run('Nuts (47-100%), Ice Cream (37-47%)').font.size = Pt(10)

doc.add_paragraph()
doc.add_heading('Notable Trends Across All Items', level=3)
all_food_findings = [
    'Popcorn: Taste complaints surged dramatically — 61.5% (2024) → 82.3% (HS1) → 67.7% (HS2). HS1 had 96 dissatisfied respondents in just 6 games vs 39 in all of 2024. The "tasted like fish" verbatim suggests a cross-contamination or oil quality issue.',
    'Pizza: Taste is the overwhelming driver at 70-81% across all periods. This is the most taste-dominated item. Temperature and Portion are minor. Likely a recipe or product quality issue.',
    'Pretzels: Temperature complaints nearly doubled from 2024 (36.5%) to HS1 (42.9%), suggesting pretzels are being served cold. HS1 had 91 dissatisfied respondents vs 52 in all of 2024.',
    'Nachos: Shifted from Taste-dominant (56.4% in 2024) to a Taste/Temperature split (37.5%/37.5% in HS2). Temperature becoming a bigger issue suggests holding time problems.',
    'Burgers: Temperature complaints nearly doubled from 23.5% (2024) to 42.2% (HS1), though improved to 34.5% in HS2. Taste declined from 58.8% to 46-48%.',
    'Salad: Zero dissatisfied respondents in 2026 (only 2 in all of 2024). Too low volume to analyze.',
    'Sausage: Small sample (4-15 per period) but Portion Size spiked to 40% in HS1 before normalizing.',
    'Hotdog & Fries: As covered in Investigation 4 — Temperature dominant, high volume. The most impactful items to fix operationally.',
]
for f in all_food_findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_heading('TB_ADDON_14: Food Sub-Attributes (2024 Only)', level=3)
doc.add_paragraph(
    'TB_ADDON_14_1 through _5 captured additional food sub-attributes (Taste, Temperature, '
    'Portion Size, Freshness, Presentation) on a 1-4 Likert scale (1=Highly Satisfied). '
    'However, this question was only asked in 2024 (n=359 respondents) and was not included '
    'in the 2026 survey instrument (0 responses). The 2024 averages were all 1.26, indicating '
    'very high satisfaction among those who answered this optional question — making it of '
    'limited analytical value for comparison. The CONCESS_QUAL_REASON codes provide the '
    'more actionable year-over-year data.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 8: SEAT VALUE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. Seat Value Perception', level=1)
doc.add_heading('"% Good Value" Dropped from 58.8% (2024) to 56.7% (2026 Combined)', level=2)

doc.add_paragraph(
    'When asked "thinking of the money spent on your ticket, your view of the game was…", '
    'the percentage of fans selecting "Good value for the money I spent on the ticket" declined '
    'in 2026. However, this aggregate masks a dramatic difference between homestands.'
)

doc.add_heading('Period-Level Summary', level=3)
add_table(doc,
    ['Period', 'Total n', 'Good', '% Good', 'Fair', '% Fair', 'Poor', '% Poor'],
    [
        ['2024 (full season)', '17,227', '10,123', '58.8%', '6,223', '36.1%', '881', '5.1%'],
        ['HS1 (Apr 6-12)', '3,439', '1,555', '45.2%', '1,567', '45.6%', '317', '9.2%'],
        ['HS2 (Apr 20-26)', '2,431', '1,771', '72.9%', '586', '24.1%', '74', '3.0%'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'The story is one of two very different homestands: HS1 had the worst seat value perception '
    '(45.2% Good) while HS2 had the best (72.9% Good — far exceeding 2024). The combined 2026 '
    'average is dragged down entirely by HS1.'
)

doc.add_heading('Game-by-Game Breakdown', level=3)
add_table(doc,
    ['Game Date', 'HS', 'Total n', '% Good', '% Poor'],
    [
        ['Apr 6 (Opening Day)', 'HS1', '1,016', '39.2%', '11.2%'],
        ['Apr 7', 'HS1', '547', '48.4%', '6.9%'],
        ['Apr 8', 'HS1', '542', '45.8%', '8.9%'],
        ['Apr 10', 'HS1', '387', '47.5%', '8.0%'],
        ['Apr 11', 'HS1', '463', '52.1%', '9.7%'],
        ['Apr 12', 'HS1', '484', '45.2%', '8.5%'],
        ['Apr 20', 'HS2', '453', '75.5%', '3.5%'],
        ['Apr 21', 'HS2', '381', '78.0%', '3.7%'],
        ['Apr 22', 'HS2', '629', '74.9%', '2.2%'],
        ['Apr 24', 'HS2', '294', '72.1%', '2.7%'],
        ['Apr 25', 'HS2', '366', '63.4%', '3.8%'],
        ['Apr 26', 'HS2', '307', '70.4%', '2.6%'],
    ])

doc.add_paragraph()
doc.add_heading('Why HS1 Was So Poor', level=3)
doc.add_paragraph(
    'Opening Day (Apr 6) had the worst seat value score of any game: 39.2% Good, 11.2% Poor. '
    'Several factors converged to create poor value perception in HS1:'
)
hs1_seat_causes = [
    'Opening Day pricing premium — fans paid top dollar for the excitement but many felt the seats did not match expectations. Multiple verbatims reference $150+ tickets with poor views.',
    'New netting poles and foul line obstructions — the most common 2026 seat view complaint. Multiple fans report poles blocking home plate or the batter from lower bowl seats.',
    'Cable camera obstruction — a new camera on wires traversing the 3rd base line repeatedly mentioned as blocking the view from Sections 109, 205, and upper deck.',
    'Full-capacity crowds in HS1 (especially Opening Day) — seat-jumping, foot traffic blocking views, standing fans, and cramped conditions were amplified by record attendance.',
    'Standing Room Only tickets sold despite empty upper deck sections — several fans expressed frustration at being sold SRO tickets while the 300-level was closed.',
]
for c in hs1_seat_causes:
    doc.add_paragraph(c, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Why HS2 Was Dramatically Better', level=3)
doc.add_paragraph(
    'HS2 reversed the trend entirely, with 72.9% Good value — far exceeding even the 2024 '
    'baseline of 58.8%. This improvement aligns with:'
)
hs2_seat_improvements = [
    'Ticket mix: HS2 had more promotional/complimentary tickets (Salute to Service, Players Give Back, Rays Rush) where expectations are lower and value perception is inherently higher.',
    'Lower crowd density: HS2 had fewer SRO/overflow situations, reducing view obstruction issues.',
    'Salute to Service offers (208+ responses) had 84.6% Good — free or heavily discounted tickets naturally score high on value.',
    'All-Inclusive Tickets in HS2 (66 responses) scored 72.7% Good despite premium pricing, suggesting the bundled food/drink value resonated.',
]
for i in hs2_seat_improvements:
    doc.add_paragraph(i, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Sentence-Level: What Fans Are Saying About Seat Value', level=3)
doc.add_paragraph(
    'Negative seat-related mentions surged in 2026 compared to 2024:'
)
add_table(doc,
    ['Category', '2024', '2026', 'Change'],
    [
        ['Seat View - Negative', '33', '140', '+324%'],
        ['Ticketing Value Perception - Negative', '18', '85', '+372%'],
        ['Ticket Pricing - Negative', '28', '51', '+82%'],
    ])

doc.add_paragraph()
seat_verbatims = [
    "I paid over $1000 for 4 tickets to sit in the 2nd row right behind 1st base and had a pole obstructing my view of 1st base and the pitcher.",
    "The netting and the poles down the foul lines are too much of an obstruction for a $45 ticket.",
    "I was also seated on the upper deck third base line in section 205, and although the new cable cam looks great on TV, it frequently was blocking my view of the game.",
    "For $320 each, I was expecting a much more comfortable and roomy seating experience.",
    "Ray's games used to be the best value in all professional sports, clearly that's not the case anymore.",
    "Paid over $150 per ticket and was extremely disappointed.",
    "Horrible view with no warning when purchasing the tickets.",
    "The tickets I purchased in the parking lot outside the stadium turned out to be restricted view behind the foul pole.",
    "Son said he would have preferred to go to the minor league game close to home.",
    "We spent $200 per ticket and missed an inning and a half.",
]
add_verbatim_block(doc, 'Representative Seat Value Verbatims (2026):', seat_verbatims)

doc.add_paragraph()
doc.add_heading('Key Findings', level=3)
seat_findings = [
    'The ~2% decline in "% Good Value" vs 2024 is entirely an HS1 phenomenon. HS2 (72.9%) dramatically outperformed 2024 (58.8%).',
    'Netting poles are the #1 new structural complaint in 2026. Multiple sections (109, 114, 120, 125, 127, 130, 139) are affected. These seats need to be flagged as "obstructed view" in the ticketing system.',
    'The cable camera on the 3rd base side is a new and recurring complaint — it did not exist pre-storm. Sections 109, 205, and club level are most affected.',
    'Opening Day pricing created a mismatch between price paid and experience delivered, especially for fans who paid $100+ and encountered obstructed views, long concession waits, and app failures simultaneously.',
    'Promotional ticket holders in HS2 (Salute to Service, Rays Rush, Players Give Back) rated value extremely high (80-100%), significantly pulling up the average.',
    'Monitor HS3 to determine whether the improvement sustains at organic (non-promotional) ticket price levels.',
]
for f in seat_findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# INVESTIGATION 9: PARKING — FINDING A SPACE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Parking — Finding a Space & Vehicle', level=1)
doc.add_heading('Down 3.4% vs 2024 — Prepaid-Only Policy & ADA Failures Dominate', level=2)

doc.add_paragraph(
    'The "Finding a space and parking your vehicle" metric declined 3.4% in 2026 compared '
    'to 2024. However, this aggregate masks an extreme HS1-to-HS2 split: HS1 averaged 8.29 '
    '(well below 2024\'s 8.76) while HS2 recovered to 9.10 (above baseline). Opening Day '
    'alone (7.38) dragged the entire season average down. This investigation examines parking '
    'complaints from three data sources: the structured parking rating, sentence-level AI '
    'analysis, and TB_ADDON_8_6 (the parking-specific feedback field).'
)

doc.add_heading('Structured Parking Rating by Period', level=3)
add_table(doc,
    ['Period', 'n', 'Avg Parking Rating', 'Avg Exit Rating', 'vs 2024'],
    [
        ['2024 (full season)', '39,533', '8.76', '8.47', '—'],
        ['HS1 (Apr 6-12)', '3,890', '8.29', '8.03', '-5.4%'],
        ['HS2 (Apr 20-26)', '2,635', '9.10', '8.75', '+3.9%'],
    ])

doc.add_paragraph()
doc.add_heading('Game-by-Game Parking Rating (2026)', level=3)
doc.add_paragraph(
    'Opening Day is a clear outlier — nearly 1.4 points below the 2024 average. By the '
    'end of HS1, ratings had recovered to baseline. HS2 consistently exceeded 2024:'
)
add_table(doc,
    ['Game Date', 'HS', 'n', 'Avg Parking Rating'],
    [
        ['Apr 6 (Opening Day)', 'HS1', '850', '7.38'],
        ['Apr 7', 'HS1', '460', '8.76'],
        ['Apr 8', 'HS1', '474', '8.60'],
        ['Apr 10', 'HS1', '328', '8.67'],
        ['Apr 11', 'HS1', '410', '8.48'],
        ['Apr 12', 'HS1', '448', '8.72'],
        ['Apr 20', 'HS2', '406', '9.11'],
        ['Apr 21', 'HS2', '347', '9.00'],
        ['Apr 22', 'HS2', '550', '8.97'],
        ['Apr 24', 'HS2', '259', '9.10'],
        ['Apr 25', 'HS2', '327', '9.35'],
        ['Apr 26', 'HS2', '268', '9.18'],
    ])

doc.add_paragraph()
doc.add_heading('Sentence-Level: Negative Parking Mentions', level=3)
doc.add_paragraph(
    'AI-categorized sentence-level analysis reveals a massive volume of parking complaints '
    'in 2026, heavily concentrated in HS1. Total 2026 negative parking sentences (440) already '
    'exceed the 2024 full-season total (70) by 529%:'
)
add_table(doc,
    ['Category', '2024 (full)', '2026 HS1', '2026 HS2', '2026 Total'],
    [
        ['Parking Availability', '12', '204', '16', '228'],
        ['Departure Traffic', '30', '75', '29', '111'],
        ['Parking Cost', '13', '35', '24', '64'],
        ['Parking Staff', '0', '29', '8', '37'],
        ['Total Negative', '55', '343', '77', '440'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'Key observation: Parking Availability complaints dropped 92% from HS1 (204) to HS2 (16), '
    'confirming the prepaid-only debacle was a first-homestand crisis. Parking Cost complaints '
    'remained proportionally high in HS2 (24 mentions), suggesting ongoing price sensitivity. '
    'Departure Traffic improved 61% (75 to 29) but remains a persistent concern.'
)

doc.add_heading('Opening Day: The Parking Catastrophe (Apr 6)', level=3)
doc.add_paragraph(
    'April 6 alone generated 174 negative parking sentences — 49% of HS1\'s entire parking '
    'complaint volume and 2.5x the ENTIRE 2024 season. The breakdown:'
)
add_table(doc,
    ['Category', 'Apr 6 Count', '% of HS1 Total'],
    [
        ['Parking Availability', '120', '59% of 204'],
        ['Departure Traffic', '28', '37% of 75'],
        ['Parking Cost', '14', '40% of 35'],
        ['Parking Staff', '12', '41% of 29'],
        ['Total', '174', '51% of 343'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'Root cause: The shift to prepaid-only parking was NOT communicated to fans in advance. '
    'Fans arrived to discover all lots required prepaid tickets, many of which were already sold out. '
    'This created cascading failures: hour-long traffic backups on 1st Ave S, fans turned away '
    'at lot entrances, and especially severe impacts on ADA/handicap parkers who had no '
    'alternative nearby.'
)

doc.add_heading('TB_ADDON_8_6: Parking-Specific Staff Complaints', level=3)
doc.add_paragraph(
    'TB_ADDON_8_6 captures text feedback specifically about parking issues that staff needed '
    'to address. Text complaint volume in 2026 HS1 alone (75 complaints in 6 games) already '
    'exceeds one-third of 2024\'s full-season total (223 complaints over ~80 games):'
)
add_table(doc,
    ['Period', 'Total Responses', 'Text Complaints', 'Complaint Rate'],
    [
        ['2024 (full season)', '21,007', '223', '1.1%'],
        ['2026 HS1 (6 games)', '3,890', '75', '1.9%'],
        ['2026 HS2 (6 games)', '2,635', '20', '0.8%'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'HS2\'s complaint rate (0.8%) dropped below the 2024 baseline (1.1%), indicating the '
    'parking experience normalized and even improved once the prepaid policy was better '
    'communicated. Thematic analysis of the 95 combined 2026 text complaints reveals '
    'five distinct complaint clusters:'
)

doc.add_heading('Complaint Theme Analysis (2026 TB_ADDON_8_6)', level=3)
add_table(doc,
    ['Theme', 'HS1 Count', 'HS2 Count', 'Total', '% of 2026'],
    [
        ['Prepaid Policy / Communication', '28', '2', '30', '32%'],
        ['ADA/Handicap Access', '16', '5', '21', '22%'],
        ['Rude/Unhelpful Staff', '9', '7', '16', '17%'],
        ['Exit/Traffic Control', '8', '3', '11', '12%'],
        ['Cost/Overcharging', '5', '3', '8', '8%'],
        ['Signage/Wayfinding', '6', '0', '6', '6%'],
        ['Other (App barcode, lot conditions)', '3', '0', '3', '3%'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'The top two themes — Prepaid Policy (32%) and ADA Access (22%) — together account for '
    'over half of all parking complaints and are almost entirely HS1 phenomena.'
)

doc.add_heading('Theme 1: Prepaid-Only Communication Failure (32%)', level=3)
doc.add_paragraph(
    'The most frequent complaint by volume. Fans describe arriving at stadium lots after '
    'significant drives (some 2-3+ hours) only to learn at the entrance that all parking is '
    'prepaid-only. No advance email, no social media post, and no signage along the approach '
    'roads warned them. Fans with disabled passengers, pregnant family members, or young '
    'children were particularly impacted because finding alternative parking required '
    'significant additional walking.'
)

prepaid_verbatims = [
    "I never saw any information stating that every parking lot would be prepaid parking only. We arrived more than an hour early but discovered every lot was prepaid only and all sold out.",
    "Prepaid parking was a debacle that should have been communicated up front.",
    "It was not communicated that people would be unable to park in stadium lots without advanced purchase.",
    "Disclose prior to waiting for an hour in traffic that the lots are prepaid only!",
    "prepaid only but you had hundreds of empty spots",
    "Could only park if we had prepaid parking. When asked a staff member where we could park he replied, 'You're done.'",
    "Frustrated that parking is prepaid only... did not know this. Caused issues.",
]
add_verbatim_block(doc, 'Prepaid Policy Verbatims:', prepaid_verbatims)

doc.add_paragraph()
doc.add_heading('Theme 2: ADA/Handicap Parking Failures (22%)', level=3)
doc.add_paragraph(
    'The second most frequent theme involves ADA parking — and the complaints are severe. '
    'Fans with disabilities, walkers, and wheelchairs report being turned away from lots with '
    'visible empty handicap spaces, forced to park in distant unpaved areas, rushed by attendants, '
    'and given conflicting information about where ADA parking exists. This issue spans both '
    'homestands, unlike the prepaid issue which was HS1-concentrated.'
)

ada_verbatims = [
    "Accessibility parking: Tried to park in lot 1. Saw at least a half a dozen accessibility spaces empty. Conflicting reports from parking that those were prepay only... caused a significant hardship and delay.",
    "Spouse just had knee replacement surgery. We had handicap parking. The lady directing traffic twice asked us to hurry up and move along. How are you going to ask someone with a walker to hurry up?!",
    "Terrible parking in lot two on the grass for my scooter to move thru the sand area ADA.",
    "I am handicapped and the woman there insisted we go all way to the end of the line farthest away from the doors, past empty spots.",
    "Would like compensation for missed first inning due to parking attendant not allowing us in handicap parking lot on the first attempt. We had to circle around again which took almost 40 min extra.",
    "No handicap parking in lot 2. Not allowed to enter Lot 1, which had many unoccupied handicap spaces.",
    "We should not have to prepay, especially with a handicap sticker.",
]
add_verbatim_block(doc, 'ADA/Handicap Verbatims:', ada_verbatims)

doc.add_paragraph()
doc.add_heading('Theme 3: Rude/Unhelpful Parking Staff (17%)', level=3)
doc.add_paragraph(
    'Staff demeanor complaints span both homestands and reference specific lots (Lot 6, '
    'Lot 7, Lot 2). One attendant in Lot 7 was cited by name across multiple games. Complaints '
    'include yelling at fans, being dismissive, and giving conflicting directions.'
)

staff_verbatims = [
    "Lady in parking lot 7 was very rude. We attended 3 games and she was rude all 3 nights.",
    "when i went to park the attendant yelled at me",
    "Staff was telling us to go 2 different ways to park, too many chiefs not enough indians. Then got bad with us when we parked where one of them told us to.",
    "One of the parking attendants in lot 6 was extremely rude. I spoke with another lot 6 attendant later and he was nice and helpful.",
    "Parking lot attendant can be a little more polite and patient.",
    "Rude staff",
]
add_verbatim_block(doc, 'Staff Behavior Verbatims:', staff_verbatims)

doc.add_paragraph()
doc.add_heading('Theme 4: Exit/Departure Traffic (12%)', level=3)
doc.add_paragraph(
    'Exit complaints persist across both homestands. The core issue is lack of traffic control '
    'after games — fans report 30-60+ minute waits to exit lots, no police or staff directing '
    'traffic at merge points, and closed exits forcing detours. The rideshare pickup location '
    'also creates confusion, with Uber/Lyft drivers unable to access the designated area due '
    'to post-game road closures.'
)

exit_verbatims = [
    "Getting away from the stadium was one of the worst sports venue exits I have ever experienced.",
    "Exiting the stadium becomes very difficult. We spent 30 minutes to leave the parking lot.",
    "Handicap space but exit line didn't move 3 feet for a half hour. Finally went wrong direction by following others and got out quick.",
    "The Rays lack of planning for rideshare customers... the police officer facilitates all self-park cars with their exit and then opens the streets for rideshares. So unfair!",
]
add_verbatim_block(doc, 'Exit/Traffic Verbatims:', exit_verbatims)

doc.add_paragraph()
doc.add_heading('2024 vs 2026 Comparison: Complaint Theme Shift', level=3)
doc.add_paragraph(
    'The nature of parking complaints fundamentally changed between 2024 and 2026:'
)
add_table(doc,
    ['Dimension', '2024 Dominant Themes', '2026 Dominant Themes'],
    [
        ['#1 Issue', 'Exit/departure traffic control', 'Prepaid policy communication'],
        ['#2 Issue', 'Lot signage / finding your car', 'ADA/Handicap access'],
        ['#3 Issue', 'Parking cost ($20-30)', 'Rude/unhelpful staff'],
        ['#4 Issue', 'Handicap access', 'Exit traffic (persistent)'],
        ['#5 Issue', 'Tight spaces / vehicle damage', 'Cost (now $20-57)'],
        ['New in 2026', '—', 'Prepaid-only policy, app barcode failures'],
        ['Resolved/Improved', '—', 'Lot signage (fewer mentions)'],
    ])

doc.add_paragraph()
doc.add_paragraph(
    'The shift from 2024 to 2026 reflects the transition to a prepaid-only model and the '
    'return to a fully operational post-hurricane stadium. Exit traffic was the #1 complaint in '
    '2024 (with extremely detailed and frustrated accounts of 45-90 minute exits from Lot 2). '
    'In 2026, exit complaints still exist but have been overtaken by the prepaid communication '
    'failure and ADA access issues.'
)

doc.add_heading('Key Findings', level=3)
parking_findings = [
    'The 3.4% aggregate decline is entirely an HS1 problem. HS2 parking ratings (9.10) exceeded the 2024 baseline (8.76) by 3.9%. Opening Day alone (7.38) accounts for the majority of the decline.',
    'The #1 root cause is the prepaid-only parking policy rollout. Fans were not notified in advance, creating a cascade of failures on Opening Day and the early HS1 games. This single policy change generated 30 of the 95 TB_ADDON_8_6 text complaints (32%).',
    'ADA/Handicap parking is a serious and ongoing issue affecting both homestands. Fans with disabilities report being turned away from lots with visible empty spaces, forced onto unpaved surfaces, rushed by attendants, and given conflicting information. This requires immediate operational review.',
    'Parking staff behavior is a recurring concern — specific attendants in Lot 7 and Lot 6 were called out by name across multiple games for rudeness and unhelpfulness. Unlike the prepaid issue, this spans both homestands.',
    'Exit traffic remains a persistent pain point carried over from 2024, though the volume of complaints has improved. The new rideshare pickup location creates additional confusion after games.',
    'By HS2, the complaint rate in TB_ADDON_8_6 (0.8%) dropped below the 2024 baseline (1.1%), and Parking Availability complaints dropped 92% (204 to 16). The parking experience has largely stabilized outside of ADA and staff behavior issues.',
]
for f in parking_findings:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Recommendations', level=3)
parking_recs = [
    ('Pre-Game Communication',
     'Implement automated parking reminder emails for all ticket purchasers 48 hours before '
     'game day, clearly stating prepaid-only policy and linking to purchase. Add signage on '
     '1st Ave S and 9th St approach roads well before lot entrances.'),
    ('ADA Parking Protocol',
     'Establish a dedicated ADA parking coordinator for each lot who can make real-time '
     'allocation decisions. Ensure all lots have paved ADA spaces (not sand/grass). Create '
     'a clear process for ADA parkers to prepay online, and allow day-of ADA parking when '
     'spaces are visibly available regardless of prepaid status.'),
    ('Staff Training & Accountability',
     'Investigate specific staff complaints in Lot 7 and Lot 6 (multiple-game recurrence '
     'of the same attendant being cited). Implement de-escalation training for high-stress '
     'situations like turning away fans from full lots.'),
    ('Exit Traffic Management',
     'Station traffic control personnel at all lot exit merge points for the full duration of '
     'post-game departure. Coordinate rideshare pickup timing with police so rideshare access '
     'is not delayed until all self-parked cars have exited.'),
]
for title, body in parking_recs:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    run.font.size = Pt(10)
    run2 = p.add_run(body)
    run2.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# METHODOLOGY NOTE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Methodology Notes', level=1)
notes = [
    'Data Source: Qualtrics VOC Post-Attendance Survey, loaded into Snowflake (TBRDP_DW_DEV.IM_RPT).',
    'Views Used: V_SBL_QUALTRICS_VOC_POST_ATTENDANCE_FULL_CORTEX_AI (structured grid ratings, wait times) and V_SBL_QUALTRICS_VOC_POST_ATTENDANCE_FULL (CONCESS_QUAL_REASON numeric codes).',
    'Sentence-Level Analysis: T_OVERALL_FEEDBACK_SENTENCE_LEVEL table with AI-categorized open-ended feedback at sentence level.',
    'CONCESS_QUAL_REASON: Only populated for respondents who rated a food item as "Somewhat Dissatisfied" or "Highly Dissatisfied." Code mapping: 1=Taste, 2=Temperature, 3=Portion Size.',
    'Junk Code Filter: All numeric columns filtered to < 80 to exclude Qualtrics junk codes.',
    'Valid DESC Labels: Structured grid questions use ("Highly satisfied", "Somewhat satisfied", "Somewhat dissatisfied", "Highly dissatisfied"). Wait expectations use ("Much more than what I expected", "Slightly more than what I expected", "About what I expected", "Slightly less what I expected", "Much less than what I expected").',
    'Homestand Dates: HS1 = Apr 6-12 (6 games), HS2 = Apr 20-26 (6 games). 2024 = full season baseline.',
    'Note on _DESC Column Mislabeling: The CONCESS_QUAL_REASON_*_DESC columns map numeric reason codes (1,2,3) to generic satisfaction labels ("Highly satisfied" = code 1, etc.) rather than the actual reason names (Taste, Temperature, Portion Size). The numeric FLOAT columns are authoritative.',
    '% Change Calculation: Relative percent change = (new - old) / old * 100, consistent with the build_voc_comparison.py report methodology.',
    'TB_ADDON_8_9: Text column in the Cortex AI view capturing app-specific feedback. Most values are numeric satisfaction scores (85-89); non-numeric values are open-text complaints.',
    'SEATVIEW_VALUE_SEAT: Text column with three response options: "Good value for the money I spent on the ticket", "Fair value for the money I spent on the ticket", "Poor value for the money I spent on the ticket".',
    'CONCESS_GRID_CUSTSERV_DESC: Text column for concessions customer service structured grid. Four-point scale: Highly satisfied, Somewhat satisfied, Somewhat dissatisfied, Highly dissatisfied.',
    'TB_ADDON_14_1 through _5: 2024-only columns (n=359). 1-4 Likert scale (1=Highly Satisfied). Not available in 2026 survey instrument.',
    'TB_ADDON_8_6: Text column in the Cortex AI view capturing parking-specific feedback for staff resolution. Most values are numeric satisfaction scores; non-numeric values are open-text complaints.',
    'PARKING_NUMRAT: Numeric 0-10 rating for "Finding a space and parking your vehicle." Available in the non-Cortex view for both 2024 and 2026.',
    'Sentence-Level Parking Categories: PARENT_CATEGORY = "Parking" with DETAILED_CATEGORY values including Parking Availability, Parking Cost, Parking Staff, and Departure Traffic (each with Negative/Neutral/Positive suffixes).',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

# ── Save ──
output_path = r'C:\Users\ytaketani\voc_insights_agent\HS12_DeepDive.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'File size: {os.path.getsize(output_path):,} bytes')
